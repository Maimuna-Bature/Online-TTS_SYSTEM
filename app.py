from flask import Flask, render_template, request, send_file, url_for, redirect, session, flash
import os
import re
import time
import tempfile
import asyncio
from pathlib import Path

# third-party libs used at runtime (import inside try/except where optional)
try:
    import edge_tts
except Exception:
    edge_tts = None

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'default_secret_key')

# temp dir for uploads and generated audio
mAUDIO_UPLOAD_TEMP_DIR = os.path.join(tempfile.gettempdir(), 'mtts_app_temp_files')
os.makedirs(mAUDIO_UPLOAD_TEMP_DIR, exist_ok=True)

# voices map (label -> edge_tts voice id)
EDGE_TTS_VOICES = {
    "English (US) - Aria Female": "en-US-AriaNeural",
    "English (US) - Guy Male": "en-US-GuyNeural",
    "English (UK) - Sonia Female": "en-GB-SoniaNeural",
    "English (UK) - Libby Female": "en-GB-LibbyNeural",
    "English (UK) - Ryan Male": "en-GB-RyanNeural"
}

def cleanup_temp_files(max_age_seconds: int = 3600):
    now = time.time()
    for name in os.listdir(mAUDIO_UPLOAD_TEMP_DIR):
        path = os.path.join(mAUDIO_UPLOAD_TEMP_DIR, name)
        try:
            if os.path.isfile(path) and (now - os.path.getmtime(path) > max_age_seconds):
                os.remove(path)
        except Exception:
            pass

cleanup_temp_files()

async def synthesize_edge_tts_async(text: str, voice: str, output_path: str):
    if edge_tts is None:
        raise RuntimeError("edge_tts package is not available.")
    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(output_path)

def process_text_for_speech(text: str, is_image: bool = False) -> str:
    if not text:
        return ""
    # basic normalization and cleanup
    text = text.replace('\r', ' ')
    text = re.sub(r'\n{2,}', '\n', text)
    text = ' '.join(text.split())
    text = re.sub(r'_{3,}', '', text)
    text = re.sub(r'[\-_]{2,}', '', text)
    text = text.replace('*', ', ')
    text = text.replace('•', ', ')
    text = text.replace('...', '…')
    # ensure spacing after punctuation
    text = re.sub(r'([.?!,:;])([^\s])', r'\1 \2', text)
    return text.strip()

@app.route('/tts', methods=['GET', 'POST'])
def index():
    error_message = None
    processed_text = ""
    audio_filename = None
    custom_filename = ""
    selected_voice = list(EDGE_TTS_VOICES.values())[0]
    success_message = None
    file_name = None
    file_size_display = None

    if request.method == 'POST':
        text_input = (request.form.get('text') or '').strip()
        uploaded_file = request.files.get('file')
        selected_voice = request.form.get('voice', selected_voice)

        # handle upload if present
        if uploaded_file and uploaded_file.filename:
            file_name = uploaded_file.filename
            tmp_path = os.path.join(mAUDIO_UPLOAD_TEMP_DIR, file_name)
            try:
                uploaded_file.save(tmp_path)
                file_size_display = f"{round(os.path.getsize(tmp_path)/1024, 2)} KB"
                ext = Path(file_name).suffix.lower()
                # DOCX
                if ext == '.docx':
                    try:
                        import docx
                        doc = docx.Document(tmp_path)
                        parts = []
                        for p in doc.paragraphs:
                            if p.text.strip():
                                parts.append(p.text)
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                                if row_text:
                                    parts.append(row_text)
                        raw = '\n'.join(parts)
                        processed_text = process_text_for_speech(raw)
                    except Exception as e:
                        error_message = f"Error reading .docx: {e}"
                # PDF
                elif ext == '.pdf':
                    try:
                        from pdfminer.high_level import extract_text
                        raw = extract_text(tmp_path)
                        processed_text = process_text_for_speech(raw)
                    except Exception as e:
                        error_message = f"Error reading PDF: {e}"
                # Image: OCR
                elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
                    try:
                        import pytesseract
                        from PIL import Image, ImageOps, ImageFilter
                        # optional tesseract path from env
                        t_cmd = os.environ.get('TESSERACT_CMD')
                        if t_cmd and os.path.exists(t_cmd):
                            pytesseract.pytesseract.tesseract_cmd = t_cmd

                        img = Image.open(tmp_path)
                        if img.mode != 'RGB':
                            img = img.convert('RGB')

                        # upscale small images
                        if img.width < 1000:
                            scale = max(2, int(1000 / max(1, img.width)))
                            img = img.resize((img.width * scale, img.height * scale), Image.LANCZOS)

                        gray = ImageOps.grayscale(img)
                        gray = gray.filter(ImageFilter.MedianFilter(size=3))
                        gray = ImageOps.autocontrast(gray, cutoff=2)

                        # attempt adaptive threshold using OpenCV if available
                        try:
                            import cv2
                            import numpy as np
                            arr = np.array(gray)
                            th = cv2.adaptiveThreshold(arr, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                                       cv2.THRESH_BINARY, 41, 15)
                            proc = Image.fromarray(th)
                        except Exception:
                            proc = gray.point(lambda p: 255 if p > 160 else 0)

                        configs = ["--oem 3 --psm 6", "--oem 3 --psm 3", "--oem 3 --psm 11"]
                        extracted = ""
                        for cfg in configs:
                            try:
                                txt = pytesseract.image_to_string(proc, lang='eng', config=cfg)
                                if txt and txt.strip():
                                    extracted = txt.strip()
                                    break
                            except Exception:
                                continue
                        if not extracted:
                            try:
                                extracted = pytesseract.image_to_string(gray, lang='eng', config="--oem 3 --psm 3").strip()
                            except Exception:
                                extracted = ""

                        if extracted:
                            processed_text = process_text_for_speech(extracted, is_image=True)
                        else:
                            error_message = "No readable text was found in the image. Try a clearer image."
                    except Exception as e:
                        error_message = f"OCR error: {e}"
                else:
                    error_message = "Unsupported file type. Use .docx, .pdf, .jpg, .png."
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

        # if no upload, use textarea
        elif text_input:
            processed_text = process_text_for_speech(text_input)
        else:
            error_message = "Please provide text input or upload a supported file."

        # perform TTS if we have text and no error
        if not error_message and processed_text.strip():
            if edge_tts is None:
                error_message = "TTS engine not available (edge_tts missing)."
            else:
                try:
                    audio_filename = f"m{int(time.time())}_{os.getpid()}.mp3"
                    audio_path = os.path.join(mAUDIO_UPLOAD_TEMP_DIR, audio_filename)
                    # use edge_tts directly (async)
                    asyncio.run(synthesize_edge_tts_async(processed_text, selected_voice, audio_path))
                    session['audio_filename'] = audio_filename
                    session['selected_voice'] = selected_voice
                    session['success_message'] = "✅ Conversion complete! Your audio is ready."
                    return redirect(url_for('index'))
                except Exception as e:
                    error_message = f"TTS conversion error: {e}"

        # if we reached here, render with messages below

    # GET or after POST fallback - pick messages from session if set
    audio_filename = session.pop('audio_filename', None)
    selected_voice = session.pop('selected_voice', list(EDGE_TTS_VOICES.values())[0])
    success_message = session.pop('success_message', None)

    return render_template(
        'index.html',
        error_message=error_message,
        processed_text=processed_text,
        audio_filename=audio_filename,
        file_name=file_name,
        file_size=file_size_display,
        edge_tts_voices=EDGE_TTS_VOICES,
        selected_voice=selected_voice,
        success_message=success_message
    )

@app.route('/download/<audio_filename>')
def download_audio(audio_filename):
    custom_name = request.args.get('name', None)
    audio_path = os.path.join(mAUDIO_UPLOAD_TEMP_DIR, audio_filename)
    if not os.path.exists(audio_path):
        return "File not found", 404
    if custom_name:
        if not custom_name.lower().endswith('.mp3'):
            custom_name += '.mp3'
        return send_file(audio_path, as_attachment=True, download_name=custom_name)
    return send_file(audio_path, as_attachment=True)

@app.route('/audio/<audio_filename>')
def serve_audio(audio_filename):
    audio_path = os.path.join(mAUDIO_UPLOAD_TEMP_DIR, audio_filename)
    if os.path.exists(audio_path):
        return send_file(audio_path, mimetype='audio/mp3')
    return "File not found", 404

@app.route('/')
@app.route('/home')
def home():
    return render_template('home.html')

if __name__ == '__main__':
    # debug True only for local development
    app.run(debug=True, port=5000)
# =====================================
# Author: Maimuna Abdulkadir Usman
# Project: Unique TTS System
# Custom Version: June-September 2025
# =====================================